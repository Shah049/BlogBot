"""
AI Article & Blog Generation System (Gradio + Groq Llama3 + DOCX export)
- Uses Groq's chat completions endpoint (Llama 3.3 70B)
- Optional contextual enrichment via Wikipedia (public data)
- Exports to .docx using python-docx
- Keeps a small session history (last 5 articles)
- Does NOT use .env; set GROQ_API_KEY directly below (or use Hugging Face Secrets)
"""

import os
import io
import re
import json
import time
import uuid
import html
from typing import List, Tuple, Dict, Optional

import requests
from bs4 import BeautifulSoup
from groq import Groq
import gradio as gr
from docx import Document
from docx.shared import Pt

# ----------------------
# CONFIGURATION (edit)
# ----------------------
# Replace the string below with your Groq API key.
# WARNING: Do NOT commit real secrets to a public repo. For HF Spaces, prefer using
# the Secrets mechanism and then set: GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
GROQ_API_KEY = "gsk_your_api_key"

# Example: If deploying on Hugging Face Spaces and you'd rather use secrets:
# GROQ_API_KEY = os.environ.get("GROQ_API_KEY")  # uncomment when using HF secrets

# Model name (Groq Llama 3.3 70B family)
GROQ_MODEL = "llama-3.3-70b-versatile"

# Number of recent articles to keep in session
MAX_HISTORY = 5

# ----------------------
# Initialize Groq client
# ----------------------
client = Groq(api_key=GROQ_API_KEY)

# ----------------------
# Utility functions
# ----------------------
def wikipedia_short_summary(query: str, max_sentences: int = 3) -> str:
    """
    Use Wikipedia's public API to fetch a short summary for the topic.
    Returns an empty string if not found.
    """
    try:
        search_url = "https://en.wikipedia.org/w/api.php"
        params = {
            "action": "query",
            "list": "search",
            "srsearch": query,
            "format": "json",
            "srlimit": 1,
        }
        r = requests.get(search_url, params=params, timeout=10)
        r.raise_for_status()
        data = r.json()
        search_results = data.get("query", {}).get("search", [])
        if not search_results:
            return ""
        page_title = search_results[0]["title"]
        # fetch extract
        params2 = {
            "action": "query",
            "prop": "extracts",
            "exintro": True,
            "explaintext": True,
            "titles": page_title,
            "format": "json",
        }
        r2 = requests.get(search_url, params=params2, timeout=10)
        r2.raise_for_status()
        data2 = r2.json()
        pages = data2.get("query", {}).get("pages", {})
        for p in pages.values():
            extract = p.get("extract", "")
            if not extract:
                return ""
            # return first N sentences
            sentences = re.split(r'(?<=[.!?]) +', extract.strip())
            return " ".join(sentences[:max_sentences])
    except Exception:
        return ""

def build_generation_prompt(
    topic: str,
    tone: str,
    include_context: bool = False,
    context_text: Optional[str] = None,
    sections: int = 5,
    target_audience: Optional[str] = None,
) -> str:
    """
    Construct a thorough system/user prompt for the Groq model to generate
    a structured article in Markdown format.
    """
    base = [
        "You are an expert content writer and SEO specialist. Produce a high-quality, "
        "well-structured article in Markdown format. Use clear H2 headings for sections (format: '## Section Title'), "
        "and include an engaging title, an SEO-optimized introduction, a multi-section body, and a concluding summary. "
        "Provide a short meta description (1-2 sentences) and a list of 5 suggested keywords at the end in JSON format like: "
        "`{\"meta\": \"...\", \"keywords\": [\"k1\",\"k2\"]}`. Keep tone consistent with the requested tone."
    ]
    base.append(f"Topic: {topic}")
    base.append(f"Tone: {tone}")
    base.append(f"Number of sections (besides intro & conclusion): {sections}")
    if target_audience:
        base.append(f"Target audience: {target_audience}")

    if include_context and context_text:
        base.append("Use the following contextual reference material to improve accuracy and relevance:")
        base.append(context_text[:4000])  # limit context size
        base.append("Incorporate this information as background — do not copy verbatim. Instead, paraphrase and cite the source as 'based on public sources' when needed.")

    instructions = (
        "Output rules:\n"
        "1) Output the article in Markdown. Use '##' for section headings. Start with a single H1 title line '# Title'.\n"
        "2) After the article, on its own lines, append a JSON object with 'meta' and 'keywords' as described.\n"
        "3) Keep the article length between 600 and 1600 words unless user requests expansion.\n"
    )
    base.append(instructions)
    return "\n\n".join(base)

def call_groq_chat(prompt: str, max_tokens: int = 2000, temperature: float = 0.2) -> str:
    """
    Call Groq chat completions. Return the assistant content as plain string.
    """
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model=GROQ_MODEL,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return response.choices[0].message.content
    except Exception as e:
        # return a useful error message
        return f"ERROR: Groq API call failed: {str(e)}"

def split_markdown_sections(markdown: str) -> Tuple[str, List[Tuple[str,str]]]:
    """
    Parse the model output markdown into title, and list of (heading, content).
    Returns: (title, sections)
    Title is the top H1 line (without '# ').
    Sections is a list ordered by appearance; headings include 'Introduction' and 'Conclusion' if present.
    """
    title = ""
    # Normalize newlines
    md = markdown.strip()
    # Extract title (H1)
    m = re.match(r"# (.+)", md)
    if m:
        title = m.group(1).strip()
        md = md[m.end():].strip()
    # Split by '## '
    parts = re.split(r"\n(?=## )", md)
    sections = []
    for part in parts:
        part = part.strip()
        if part.startswith("## "):
            lines = part.split("\n", 1)
            heading = lines[0].replace("## ", "").strip()
            content = lines[1].strip() if len(lines) > 1 else ""
            sections.append((heading, content))
        else:
            # treat as preface/introduction if no H2
            if part:
                sections.insert(0, ("Introduction", part))
    return title, sections

def create_docx_from_markdown(title: str, sections: List[Tuple[str,str]], author: Optional[str]=None) -> bytes:
    """
    Create a .docx file from parsed markdown sections and return bytes.
    """
    doc = Document()
    # Title
    h = doc.add_heading(title, level=0)
    if author:
        p = doc.add_paragraph()
        p.add_run(f"Author: {author}").italic = True
    doc.add_paragraph()  # spacing

    for heading, content in sections:
        doc.add_heading(heading, level=1)
        # content may have markdown paragraphs; simple split by double newline
        paragraphs = re.split(r'\n\s*\n', content)
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph(para)
            # keep base font reasonable
            for run in p.runs:
                run.font.size = Pt(11)

    # Add small footer meta (word count)
    full_text = " ".join([c for _, c in sections])
    words = len(re.findall(r"\w+", full_text))
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Word count: {words} · Generated by Groq Llama 3.3 via a Gradio app")
    footer.runs[0].italic = True

    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def estimate_reading_time(text: str, wpm: int = 200) -> str:
    words = len(re.findall(r"\w+", text))
    minutes = max(1, round(words / wpm))
    return f"{minutes} min read ({words} words)"

# ----------------------
# High-level features
# ----------------------
def generate_article_flow(
    topic: str,
    tone: str,
    enable_context: bool,
    sections_count: int,
    target_audience: str,
    keep_history: dict,
):
    """
    Main orchestration for generating an article:
    - optional context enrichment via Wikipedia
    - build prompt
    - call Groq
    - parse output
    - store in session history
    - return article markdown, metadata, and history
    """
    if not topic.strip():
        return "Please provide a topic or brief.", {}, keep_history

    context_text = ""
    if enable_context:
        context_text = wikipedia_short_summary(topic, max_sentences=4)
        if context_text:
            context_text = f"Wikipedia short summary: {context_text}"
        else:
            context_text = ""

    prompt = build_generation_prompt(
        topic=topic,
        tone=tone,
        include_context=bool(context_text),
        context_text=context_text,
        sections=sections_count,
        target_audience=target_audience or None,
    )

    assistant_output = call_groq_chat(prompt)
    if assistant_output.startswith("ERROR: Groq"):
        return assistant_output, {}, keep_history

    # Parse JSON meta at end if present
    meta = {}
    meta_json_match = re.search(r'\{.*"meta".*$', assistant_output, flags=re.S)
    if meta_json_match:
        # try to extract the trailing JSON by finding the first '{' from the match
        idx = assistant_output.rfind('{')
        try:
            trailing = assistant_output[idx:]
            meta = json.loads(trailing)
            # remove trailing JSON from content
            assistant_output = assistant_output[:idx].strip()
        except Exception:
            meta = {}

    title, sections = split_markdown_sections(assistant_output)
    full_text = "\n\n".join([f"## {h}\n{c}" for h,c in sections])
    reading_time = estimate_reading_time(full_text)

    # Save to history (in-memory dict storing lists)
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    article_id = str(uuid.uuid4())
    entry = {
        "id": article_id,
        "title": title or topic,
        "topic": topic,
        "tone": tone,
        "content_md": assistant_output,
        "sections": sections,
        "meta": meta,
        "reading_time": reading_time,
        "timestamp": timestamp,
    }

    hist_list = keep_history.get("articles", [])
    hist_list.insert(0, entry)
    # cap
    hist_list = hist_list[:MAX_HISTORY]
    keep_history["articles"] = hist_list

    return assistant_output, {
        "title": title or topic,
        "reading_time": reading_time,
        "meta": meta,
        "article_id": article_id,
    }, keep_history

def regenerate_section_flow(article_md: str, section_heading: str, tone: str) -> str:
    """
    Regenerate one section by instructing the model to rewrite a single section.
    Expects article_md to contain the full article in markdown.
    """
    if not section_heading:
        return "Please provide the section heading to regenerate."
    # give the model the whole article and ask to only rewrite the section
    prompt = (
        "You are an expert editor. The user will provide an article in Markdown. "
        f"Please rewrite ONLY the section titled '{section_heading}', preserving the rest of the article exactly. "
        "Keep format (use '## Section Title' heading), and keep length similar unless user asks to expand. "
        "Maintain the same tone and factual consistency. Return the full article in Markdown."
        "\n\nArticle:\n\n"
        + article_md
    )
    response = call_groq_chat(prompt)
    return response

# ----------------------
# GRADIO UI
# ----------------------
with gr.Blocks(title="AI Article & Blog Generator (Groq Llama 3.3 70B)", analytics_enabled=False) as demo:
    gr.Markdown("## 🚀 AI Article & Blog Generation System (Groq Llama 3.3 70B)\n"
                "Generate SEO-optimized blog posts, edit them, and export to DOCX. **No external image APIs**.")
    with gr.Row():
        with gr.Column(scale=2):
            topic_in = gr.Textbox(label="Topic / Brief", placeholder="e.g. Artificial Intelligence in Healthcare", lines=2)
            tone_sel = gr.Dropdown(label="Tone / Style", choices=["Professional", "Informative", "Casual", "Storytelling"], value="Professional")
            sections_count = gr.Slider(label="Number of sections (body)", minimum=2, maximum=8, step=1, value=4)
            target_audience = gr.Textbox(label="Target audience (optional)", placeholder="e.g. healthcare professionals, marketers", lines=1)
            enable_context = gr.Checkbox(label="Enrich with public context (Wikipedia)", value=True)
            generate_btn = gr.Button("Generate Article", variant="primary")
            # Hidden state for history
            history_state = gr.State({"articles": []})
        with gr.Column(scale=3):
            output_md = gr.Textbox(label="Generated Article (Markdown)", lines=20)
            info_box = gr.Markdown("", elem_id="info_box")
            with gr.Row():
                download_docx_btn = gr.Button("Download as DOCX")
                regenerate_heading = gr.Textbox(label="Section heading to regenerate (exact text from heading)", placeholder="e.g. Challenges", lines=1)
                regen_btn = gr.Button("Regenerate Section")
            with gr.Accordion("Article History (last 5)", open=False):
                history_dropdown = gr.Dropdown(label="Choose from history", choices=[], value=None)
                load_history_btn = gr.Button("Load Selected Article")

    # ----------------
    # Callbacks
    # ----------------
    def on_generate(topic, tone, enable_context, sections_count, target_audience, history):
        md, meta, updated_history = generate_article_flow(
            topic=topic,
            tone=tone,
            enable_context=enable_context,
            sections_count=int(sections_count),
            target_audience=target_audience,
            keep_history=history or {}
        )
        info = ""
        if isinstance(meta, dict) and meta.get("reading_time"):
            info = f"*Reading time:* {meta.get('reading_time')}"
        else:
            # parse title & reading time from returned meta dict (we returned them in meta earlier)
            if meta.get("title"):
                info = f"**Title:** {meta.get('title')}  ·  **Reading time:** {meta.get('reading_time')}"
            else:
                info = ""
        # update history dropdown choices
        hist_list = updated_history.get("articles", [])
        choices = [f"{i+1}. {h['title']} ({h['timestamp']})" for i, h in enumerate(hist_list)]
        # store mapping from choice to article id in the state
        updated_history["_choice_map"] = {choices[i]: hist_list[i]['id'] for i in range(len(choices))}
        return md, f"Generated · {meta.get('title','')}\n\nReading time: {meta.get('reading_time','')}", updated_history, gr.update(choices=choices)

    generate_btn.click(
        on_generate,
        inputs=[topic_in, tone_sel, enable_context, sections_count, target_audience, history_state],
        outputs=[output_md, info_box, history_state, history_dropdown],
        queue=True,
    )

    def on_regenerate(article_md, heading, tone):
        if not article_md.strip():
            return "No article loaded to regenerate.", gr.update()
        new_md = regenerate_section_flow(article_md, heading, tone)
        # If model returns error, return it directly
        if new_md.startswith("ERROR: Groq"):
            return new_md, gr.update()
        # Update history? For simplicity, we return regenerated MD but do not auto-add to history
        return new_md, gr.update()

    regen_btn.click(
        on_regenerate,
        inputs=[output_md, regenerate_heading, tone_sel],
        outputs=[output_md, info_box],
        queue=True
    )

    def on_create_docx(article_md, history):
        """
        Convert the current article to a DOCX file and prepare it for download.
        """
        if not article_md.strip():
            return None, "No article to export."

        # Parse article
        title, sections = split_markdown_sections(article_md)
        if not title:
            title = "Generated_Article"

        # Create DOCX file in memory
        file_bytes = create_docx_from_markdown(title, sections, author=None)

        # ✅ Fix: precompute sanitized filename (no backslashes in f-string)
        safe_title = re.sub(r"\W+", "_", title)[:80]
        filename = f"{safe_title}.docx"

        # Create temporary path for Hugging Face or local environments
        tmp_path = f"/tmp/{uuid.uuid4()}_{filename}"
        with open(tmp_path, "wb") as f:
            f.write(file_bytes)

        return tmp_path, f"✅ DOCX ready: {filename}"

    download_docx_btn.click(
        on_create_docx,
        inputs=[output_md, history_state],
        outputs=[gr.File(label="Download DOCX"), info_box],
        queue=True
    )

    def on_load_history(choice_label, history):
        if not choice_label:
            return "", gr.update()
        choice_map = history.get("_choice_map", {})
        article_id = choice_map.get(choice_label)
        if not article_id:
            return "Could not find article.", gr.update()
        for art in history.get("articles", []):
            if art["id"] == article_id:
                return art["content_md"], gr.update()
        return "Article not found in history.", gr.update()

    load_history_btn.click(
        on_load_history,
        inputs=[history_dropdown, history_state],
        outputs=[output_md, info_box],
    )

    # Footer
    gr.Markdown("**Notes:** This app uses Groq's Llama 3.3 70B via the Groq API for content generation. "
                "If deploying publicly, configure secrets and do not hardcode API keys in source code.")

# Expose a simple launch helper for local run or HF
def run():
    demo.launch(server_name="0.0.0.0", server_port=7860, share=False)

if __name__ == "__main__":
    run()
